import streamlit as st
import re, io, zipfile, unicodedata, subprocess, tempfile
from pathlib import Path
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter
from docx import Document

# ── Configurazione pagina ────────────────────────────────────────────────────
st.set_page_config(
    page_title="Cartelle Sanitarie",
    page_icon="📁",
    layout="centered",
)

st.title("📁 Generazione Cartelle Sanitarie")
st.markdown(
    "Carica i file TOTALI (`.docx`) e le anagrafiche (`.xlsx`). "
    "L'applicazione genera automaticamente un PDF per ogni lavoratore, "
    "organizzato in cartelle nominate **COGNOME_NOME_CF**."
)
st.info(
    "🔒 I file vengono elaborati interamente in memoria e non vengono "
    "salvati né trasmessi a terzi. Appena chiudi la pagina non rimane nulla.",
    icon="🔒",
)

# ════════════════════════════════════════════════════════════════════════════
# FUNZIONI CORE  (identiche al notebook, ma lavorano su oggetti BytesIO)
# ════════════════════════════════════════════════════════════════════════════

def sanitizza(s):
    s = str(s).upper().strip()
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    s = re.sub(r"[^A-Z0-9]+", "_", s)
    return s.strip("_")


def analizza_docx(fileobj):
    """Analizza un file docx (BytesIO) e restituisce metadati."""
    doc = Document(fileobj)
    if not doc.tables:
        return None

    tipo = "idoneita"
    for p in doc.paragraphs[:5]:
        if "CARTELLA SANITARIA" in p.text.upper():
            tipo = "cartella"
            break

    n_lavoratori = 0
    n_totale     = len(doc.tables)
    azienda      = None

    for table in doc.tables:
        cell = table.rows[0].cells[0].text
        if re.search(r'NIF\)[^:]*[:–\-]\s*\d{5,}', cell):
            n_lavoratori += 1
            if azienda is None:
                m = re.search(
                    r'Azienda[^:]*:\s*\*{0,2}([^\n*]+?)\*{0,2}\s*(?:Sede|Mansione|$)',
                    cell
                )
                if m:
                    azienda = m.group(1).strip()

    return {
        "tipo":         tipo,
        "azienda":      azienda,
        "n_lavoratori": n_lavoratori,
        "n_totale":     n_totale,
    }


def analizza_xlsx(fileobj):
    """Analizza un file xlsx (BytesIO) e restituisce azienda + lista lavoratori."""
    wb = load_workbook(fileobj, read_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    lavoratori = []
    azienda    = None
    for row in rows[1:]:
        if row[0] is None:
            continue
        cognome, nome, _, _, _, _, cf, az = row[:8]
        if azienda is None:
            azienda = str(az).strip()
        nome_cartella = f"{sanitizza(cognome)}_{sanitizza(nome)}_{int(cf)}"
        lavoratori.append(nome_cartella)

    return {"azienda": azienda, "lavoratori": lavoratori}


def docx_bytes_to_pdf_bytes(docx_bytes):
    """Converte bytes di un docx in bytes PDF tramite LibreOffice."""
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        docx_path = tmp / "input.docx"
        docx_path.write_bytes(docx_bytes)

        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(tmp), str(docx_path)],
            capture_output=True, check=True
        )
        pdf_path = tmp / "input.pdf"
        if not pdf_path.exists():
            raise RuntimeError("Conversione PDF fallita.")
        return pdf_path.read_bytes()


def dividi_pdf(pdf_bytes, n_totale, lavoratori, tipo):
    """
    Divide un PDF in blocchi e restituisce un dict:
      { nome_cartella: pdf_bytes }
    """
    reader   = PdfReader(io.BytesIO(pdf_bytes))
    n_pagine = len(reader.pages)

    if n_pagine % n_totale != 0:
        raise ValueError(
            f"Il PDF ha {n_pagine} pagine ma {n_totale} blocchi — "
            f"la divisione non è intera. Verifica il file sorgente."
        )

    ppp     = n_pagine // n_totale
    risultati = {}

    for i, nome_cartella in enumerate(lavoratori):
        writer = PdfWriter()
        for p in range(i * ppp, i * ppp + ppp):
            writer.add_page(reader.pages[p])
        buf = io.BytesIO()
        writer.write(buf)
        risultati[nome_cartella] = buf.getvalue()

    return risultati, ppp


# ════════════════════════════════════════════════════════════════════════════
# INTERFACCIA
# ════════════════════════════════════════════════════════════════════════════

uploaded = st.file_uploader(
    "Trascina qui i file oppure clicca per selezionarli",
    type=["docx", "xlsx"],
    accept_multiple_files=True,
    help="Puoi caricare file di più aziende contemporaneamente.",
)

if uploaded:
    docx_files = [f for f in uploaded if f.name.lower().endswith(".docx")]
    xlsx_files = [f for f in uploaded if f.name.lower().endswith(".xlsx")]

    st.markdown(f"**File caricati:** {len(docx_files)} docx · {len(xlsx_files)} xlsx")

    # Analisi rapida per feedback immediato
    with st.expander("📋 Dettaglio file riconosciuti", expanded=False):
        for f in docx_files:
            info = analizza_docx(io.BytesIO(f.read()))
            f.seek(0)
            if info and info["n_lavoratori"] > 1:
                st.success(
                    f"✅ **{f.name}** — {info['tipo']} · "
                    f"{info['n_lavoratori']} lavoratori · {info['azienda']}"
                )
            elif info:
                st.warning(f"⏭️ **{f.name}** — file singolo (BASE), verrà ignorato")
            else:
                st.error(f"❌ **{f.name}** — non riconoscibile")

        for f in xlsx_files:
            info = analizza_xlsx(io.BytesIO(f.read()))
            f.seek(0)
            st.info(
                f"📊 **{f.name}** — {info['azienda']} · "
                f"{len(info['lavoratori'])} lavoratori"
            )

    st.divider()

    if st.button("🚀 Genera cartelle", type="primary", use_container_width=True):

        progress = st.progress(0, text="Avvio elaborazione...")
        log      = st.empty()

        try:
            # ── Analisi ──────────────────────────────────────────────────
            docx_info = []
            for f in docx_files:
                info = analizza_docx(io.BytesIO(f.read()))
                f.seek(0)
                if info and info["n_lavoratori"] > 1:
                    info["fileobj"] = f
                    docx_info.append(info)

            xlsx_map = {}
            for f in xlsx_files:
                info = analizza_xlsx(io.BytesIO(f.read()))
                f.seek(0)
                if info["azienda"]:
                    xlsx_map[sanitizza(info["azienda"])] = info

            # ── Abbinamento ───────────────────────────────────────────────
            jobs = []
            errori = []
            for d in docx_info:
                key = sanitizza(d["azienda"] or "")
                if key not in xlsx_map:
                    errori.append(
                        f"Nessuna anagrafica xlsx per l'azienda: **{d['azienda']}**"
                    )
                    continue
                xl = xlsx_map[key]
                n_reali    = min(len(xl["lavoratori"]), d["n_lavoratori"])
                lavoratori = xl["lavoratori"][:n_reali]
                jobs.append({
                    "label":      f"{d['tipo'].upper()} – {d['azienda']}",
                    "tipo":       d["tipo"],
                    "n_totale":   d["n_totale"],
                    "lavoratori": lavoratori,
                    "fileobj":    d["fileobj"],
                })

            if errori:
                for e in errori:
                    st.error(e)
                st.stop()

            # ── Elaborazione ──────────────────────────────────────────────
            # zip in memoria
            zip_buf  = io.BytesIO()
            n_totale_jobs = len(jobs)

            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for ji, job in enumerate(jobs):
                    pct_base = ji / n_totale_jobs

                    log.markdown(f"⏳ Conversione in PDF: **{job['label']}**...")
                    progress.progress(pct_base + 0.1 / n_totale_jobs,
                                      text=f"Conversione PDF ({ji+1}/{n_totale_jobs})...")

                    pdf_bytes = docx_bytes_to_pdf_bytes(job["fileobj"].read())
                    job["fileobj"].seek(0)

                    log.markdown(f"✂️ Suddivisione: **{job['label']}** "
                                 f"({len(job['lavoratori'])} lavoratori)...")
                    progress.progress(pct_base + 0.4 / n_totale_jobs,
                                      text=f"Suddivisione ({ji+1}/{n_totale_jobs})...")

                    risultati, ppp = dividi_pdf(
                        pdf_bytes, job["n_totale"],
                        job["lavoratori"], job["tipo"]
                    )

                    for nome_cartella, pdf_data in risultati.items():
                        path_zip = f"{nome_cartella}/{job['tipo']}_{nome_cartella}.pdf"
                        zf.writestr(path_zip, pdf_data)

                    progress.progress((ji + 1) / n_totale_jobs,
                                      text=f"Completato {ji+1}/{n_totale_jobs}")

            progress.progress(1.0, text="✅ Elaborazione completata!")
            log.empty()

            n_cartelle = len(set(
                p.split("/")[0]
                for p in [zi.filename for zi in
                           zipfile.ZipFile(io.BytesIO(zip_buf.getvalue())).infolist()]
            ))
            st.success(
                f"✅ **Elaborazione completata!**  \n"
                f"{n_cartelle} cartelle generate · "
                f"{len(jobs) * 1} file PDF per tipo"
            )

            zip_buf.seek(0)
            st.download_button(
                label="⬇️  Scarica cartelle_sanitarie.zip",
                data=zip_buf,
                file_name="cartelle_sanitarie.zip",
                mime="application/zip",
                type="primary",
                use_container_width=True,
            )

        except Exception as e:
            progress.empty()
            log.empty()
            st.error(f"❌ Errore: {e}")

else:
    st.markdown(
        "<br><p style='text-align:center;color:gray;'>"
        "Nessun file caricato — trascina i file qui sopra per iniziare."
        "</p>",
        unsafe_allow_html=True,
    )
